import pandas as pd
import statsmodels.api as sm
from statsmodels.formula.api import ols
from statsmodels.stats.multicomp import pairwise_tukeyhsd
from docx import Document

# -----------------------------
# Example Data
# -----------------------------
group1 = [23, 20, 22, 21, 24]
group2 = [30, 32, 29, 31, 33]
group3 = [40, 42, 39, 38, 41]

# Combine into a DataFrame
data = {
    'value': group1 + group2 + group3,
    'group': ['group1']*5 + ['group2']*5 + ['group3']*5
}
df = pd.DataFrame(data)

# -----------------------------
# Perform ANOVA
# -----------------------------
model = ols('value ~ C(group)', data=df).fit()
anova_table = sm.stats.anova_lm(model, typ=2)

# -----------------------------
# Perform Tukey HSD Test
# -----------------------------
tukey = pairwise_tukeyhsd(endog=df['value'], groups=df['group'], alpha=0.05)

# -----------------------------
# Create Word Document
# -----------------------------
doc = Document()
doc.add_heading('ANOVA and Tukey HSD Results', 0)

# Add ANOVA results
doc.add_heading('ANOVA Table:', level=1)
doc.add_paragraph(f"F-Statistic: {anova_table['F'][0]:.4f}")
doc.add_paragraph(f"P-Value: {anova_table['PR(>F)'][0]:.4f}")

# Add a space
doc.add_paragraph('')

# Add Tukey HSD results
doc.add_heading('Tukey HSD Test Results:', level=1)

# Add Tukey Summary
summary_data = tukey.summary().data
headers = summary_data[0]
rows = summary_data[1:]

# Create a table
table = doc.add_table(rows=1, cols=len(headers))
hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = str(header)

for row in rows:
    row_cells = table.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = str(item)

# Save the document
output_filename = "ANOVA_Tukey_Results.docx"
doc.save(output_filename)

print(f"\nWord document saved successfully as '{output_filename}' in the current folder!")
